#!/usr/bin/env python3
"""Update Tinashe's eTenderScraping_Cloud_Migration_Brief.docx to address the
seven gaps identified in Chelton's review. Save as v2 next to the original.
"""
import io, sys, os, shutil, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt
from copy import deepcopy

SRC = r"C:\Users\CheltonGraham\Downloads\eTenderScraping_Cloud_Migration_Brief.docx"
DST = r"C:\Users\CheltonGraham\Downloads\eTenderScraping_Cloud_Migration_Brief_v2.docx"

shutil.copy2(SRC, DST)
doc = Document(DST)


def _find_paragraph(text_starts_with):
    """Return the first paragraph whose text starts with the given string."""
    for p in doc.paragraphs:
        if p.text.strip().startswith(text_starts_with):
            return p
    return None


def _insert_paragraph_after(paragraph, text, style=None):
    """Insert a new paragraph after the given paragraph, return the new one."""
    new_p = deepcopy(paragraph._element)
    # Clear runs from the copy
    for r in new_p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
        new_p.remove(r)
    paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = doc.styles[style]
        except KeyError:
            pass
    new_para.add_run(text)
    return new_para


def _insert_after_paragraph(anchor_para, lines):
    """Insert a sequence of (style, text) tuples after anchor_para.
    Returns the last inserted paragraph so subsequent inserts can chain."""
    cur = anchor_para
    for style, text in lines:
        cur = _insert_paragraph_after(cur, text, style=style)
    return cur


# ── GAP 1 + 2 + 3: Container-migration checklist (Section 2 additions) ──
anchor = _find_paragraph("Containerising means the scraper runs on Linux.")
if anchor:
    _insert_after_paragraph(anchor, [
        ("Heading 3", "Container-migration checklist for this repository"),
        ("Normal",
         "The remediation work Phase 1 must specifically cover for the current codebase — not left as implicit \"small changes\":"),
        ("List Paragraph",
         "Replace the Windows-only chrome-win64/ folder currently in the repo with a Linux Chromium binary installed inside the Docker image (apt-get install chromium chromium-driver in the Dockerfile)."),
        ("List Paragraph",
         "Set TZ=Africa/Johannesburg as a container environment variable. Container Apps default to UTC, but the scraper's batch labelling, log timestamps and Power BI feed all assume SAST."),
        ("List Paragraph",
         "Move config surfaces out of OneDrive-path assumptions. Specifically: Websites.xlsx (the watchlist source list read by _load_watchlist() in main.py) and config.json must either be baked into the image or mounted from Azure Blob / Azure App Configuration."),
        ("List Paragraph",
         "Re-enable Selenium headless mode before cutover. The current SeleniumWatchlistScrapers.py has --headless=new commented out to allow local visual debugging; production must re-enable it."),
        ("List Paragraph",
         "Persist the per-source snapshot files (data/snapshots/{SOURCE_KEY}.json) used by the phantom-scraper fix (Nelson Mandela Bay, Buffalo City, GDoH, etc.) across job runs. Container Apps Jobs are ephemeral — snapshots must move to either a scraper_snapshots table in Azure SQL or Azure Storage Table entities. This is a hard requirement; without persistence, snapshot-diff logic reverts to \"first run\" behaviour on every batch and re-introduces phantom counts."),
        ("List Paragraph",
         "Handle the four scraper modes (eTenders / EC DPW / All but eTenders / Full Batch) explicitly. Decide between (a) four Container Apps Jobs each with its own cron and mode-specific naming, or (b) one Job that takes a mode parameter and is invoked four times by four Logic App / Power Automate triggers. Recommended: option (a) — clearer separation, easier per-mode retry/timeout tuning, no shared-state risk."),
        ("List Paragraph",
         "Fix the known logging bug where All-but-eTenders mode does not currently write to scraper.log (root logger's FileHandler is only attached by TenderScraper.setupLogging(), which is skipped in that mode). In the container, replace the FileHandler with stdout/stderr logging; Container Apps' Log Analytics integration captures both automatically."),
    ])
    print("✓ Section 2: container-migration checklist added")


# ── GAP 5: Clarify SharePoint List position in Section 3 ──
anchor = _find_paragraph("However, I would also put SharePoint List as a contender")
if anchor:
    anchor.text = (
        "Note on SharePoint List positioning: SharePoint List is NOT suitable as the primary "
        "data store for per-tender rows (see Table 4). However, it may be worth introducing "
        "in Phase 3 as a supplementary, user-facing view — auto-populated from Azure SQL via "
        "a Power Automate flow — so non-technical Amidel staff (e.g. Thandolwethu) can browse "
        "and filter tenders directly in Teams / SharePoint without needing SQL or Power BI. "
        "This is a nice-to-have layered on top of Azure SQL, not a replacement for it."
    )
    print("✓ Section 3: SharePoint List paragraph clarified")


# ── GAP 6: Resolve manual-trigger decision in Section 6 ──
anchor = _find_paragraph("If the team later wants to adjust which watchlist sites")
if anchor:
    _insert_after_paragraph(anchor, [
        ("Heading 3", "Manual trigger — resolved scope decision"),
        ("Normal",
         "Confirming the personal preference expressed in Section 6: manual triggering IS "
         "included, but not via a custom web UI. The mechanism is:"),
        ("List Paragraph",
         "Primary: Container Apps Jobs support manual start from the Azure Portal or az CLI (Phase 1 native capability, no extra build)."),
        ("List Paragraph",
         "For non-technical Amidel staff: a Power Automate flow with a Teams-channel or Outlook button that calls the job-start REST API via service principal. This lands in Phase 3 (\"operational polish\") alongside the Teams status notifications, so both non-technical trigger and status live in the same channel."),
        ("Normal",
         "This satisfies \"scheduled + on-demand\" without building or maintaining a bespoke web UI. If a UI becomes justified later (custom watchlist edits, batch parameter overrides), the underlying job-start API remains the same target — the UI would just be another caller."),
    ])
    print("✓ Section 6: manual trigger decision resolved")


# ── GAP 4 + 7 + Phase 0: Section 8 additions ──

# Insert Phase 0 BEFORE Phase 1
anchor = _find_paragraph("Phase 1 — Containerise in place")
if anchor:
    # Insert Phase 0 above Phase 1
    from docx.text.paragraph import Paragraph
    new_p_0 = deepcopy(anchor._element)
    for r in new_p_0.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
        new_p_0.remove(r)
    anchor._element.addprevious(new_p_0)
    p0 = Paragraph(new_p_0, anchor._parent)
    try:
        p0.style = doc.styles['Heading 2']
    except KeyError:
        pass
    p0.add_run("Phase 0 — Access provisioning (approx. 1 week)")

    _insert_after_paragraph(p0, [
        ("Normal",
         "Goal: unblock every other phase before writing any code. This is the phase where "
         "cloud migrations most commonly lose their first two weeks unassuming this work "
         "\"just happens\" during Phase 1."),
        ("List Paragraph",
         "Confirm or provision the Amidel Azure subscription; agree the resource group naming convention (e.g. rg-etenderscraping-prod)."),
        ("List Paragraph",
         "Create the Azure AD App Registration for the scraper and obtain admin consent for the Sites.Selected (site-scoped) Graph API permission on the relevant SharePoint site. This is the highest-latency step because it needs the tenant admin's involvement."),
        ("List Paragraph",
         "Provision the GitHub repository, agree branch protection rules, and configure OIDC federated credentials so GitHub Actions can deploy to Azure without stored secrets."),
        ("List Paragraph",
         "Confirm SharePoint site + document library ID for the tender-report output location; validate current recipients / Bheki's flow to identify anything the migration must not disturb."),
        ("List Paragraph",
         "Exit criteria: az login from a GitHub Actions runner reaches Azure; a test file upload via Graph API to the target SharePoint library succeeds end-to-end."),
    ])
    print("✓ Section 8: Phase 0 added")

# Add acceptance criteria after Phase 1 goal line
anchor = _find_paragraph("This phase alone already satisfies the core ask")
if anchor:
    _insert_after_paragraph(anchor, [
        ("Heading 3", "Phase 1 exit / acceptance criteria"),
        ("List Paragraph",
         "Three consecutive scheduled batches run to completion on Container Apps without operator intervention."),
        ("List Paragraph",
         "For each of those three batches, the cloud end-product Tender Data sheet differs from the parallel-run laptop version by no more than 2 percent of rows (measured by normalised TENDER_ID set difference). Differences above that threshold trigger investigation, not cutover."),
        ("List Paragraph",
         "All four scraper modes (eTenders, EC DPW, All but eTenders, Full Batch) execute at least once successfully in the cloud environment before cutover, not just the primary mode."),
        ("List Paragraph",
         "Snapshot files persist across job runs (Gap 2 above) and are verified against a second run of the same batch — second run must show \"0 new / N previously seen\" for NMB / Buffalo City / GDoH scrapers."),
    ])
    print("✓ Section 8: Phase 1 acceptance criteria added")

# Add acceptance criteria after Phase 2 last bullet
anchor = _find_paragraph("Re-point the .pbix to Azure SQL and validate")
if anchor:
    _insert_after_paragraph(anchor, [
        ("Heading 3", "Phase 2 exit / acceptance criteria"),
        ("List Paragraph",
         "Every page of the migrated .pbix reports totals within 1 tender per (source × batch) of the pre-migration Excel-sourced version, across at least the six most recent batches."),
        ("List Paragraph",
         "Historical backfill of master_tenders.xlsx (~35,000 rows) is complete, indexed, and query-tested from Power BI and Excel Get Data."),
        ("List Paragraph",
         "Scheduled dataset refresh in the Power BI Service runs successfully three times in a row after each Monday and Thursday batch completes."),
    ])
    print("✓ Section 8: Phase 2 acceptance criteria added")

# Add acceptance criteria after Phase 3 last bullet
anchor = _find_paragraph("Tune Azure SQL tier and Container Apps resource allocation")
if anchor:
    _insert_after_paragraph(anchor, [
        ("Heading 3", "Phase 3 exit / acceptance criteria"),
        ("List Paragraph",
         "Non-technical user (e.g. Thandolwethu) successfully triggers a manual scrape run via the Teams / Power Automate button without developer assistance."),
        ("List Paragraph",
         "Teams status-notification message appears within 5 minutes of a batch completing, with tender count and any per-source failures called out."),
        ("List Paragraph",
         "Bicep-based full teardown-and-redeploy of the environment succeeds in a staging resource group, proving the infrastructure is truly reproducible."),
    ])
    print("✓ Section 8: Phase 3 acceptance criteria added")


# ── GAP 4: Timeline realism note (add to end of Section 8 before section 9) ──
anchor = _find_paragraph("9. Estimated monthly cost")
if anchor:
    from docx.text.paragraph import Paragraph
    new_p = deepcopy(anchor._element)
    for r in new_p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
        new_p.remove(r)
    anchor._element.addprevious(new_p)
    note_p = Paragraph(new_p, anchor._parent)
    try:
        note_p.style = doc.styles['Heading 3']
    except KeyError:
        pass
    note_p.add_run("Note on timeline realism")

    _insert_after_paragraph(note_p, [
        ("Normal",
         "The 6–11 week estimate above assumes one developer's full attention across the "
         "three phases. Factoring in Amidel's parallel commitments (existing tender-report "
         "delivery, other client work), realistic calendar time to fully-operational Phase 3 "
         "is more accurately 10–16 weeks including the new Phase 0. This does not change any "
         "cost estimate — Container Apps and Azure SQL Serverless are pay-per-run — but it "
         "should shape the go-live communication to Bheki / recipients so expectations align "
         "with actual delivery."),
    ])
    print("✓ Section 8: timeline realism note added")


# Add v2 marker to the header
first_para = doc.paragraphs[0]
if first_para.text.strip().startswith("eTenderScraping"):
    for i, p in enumerate(doc.paragraphs[:10]):
        if "23 July 2026" in p.text:
            # Update the date and add v2 marker
            p.text = p.text.replace("23 July 2026", "23 July 2026 (v2 — reviewed additions)")
            print("✓ Header: v2 marker added")
            break


doc.save(DST)
print(f"\nSaved: {DST}")

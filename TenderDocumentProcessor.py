"""Download eTender PDFs and build a consolidated bid-document report."""

import logging
import os
import re
from pathlib import Path
from typing import Iterable
from urllib.parse import urlparse

import fitz
import requests
from docx import Document
from docx.shared import Pt


_REQUIREMENT_HEADINGS = re.compile(
    r"(?:documents? constituting the bid|documents? required|supporting documents?|"
    r"bid submission|preparation of bids|mandatory requirements?)",
    re.IGNORECASE,
)
_SECTION_END = re.compile(
    r"\n\s*(?:2\.[3-9]|3\.|bid prices|period of validity|closing date|evaluation of bids)",
    re.IGNORECASE,
)


def _safe_filename(value: str) -> str:
    value = re.sub(r"[^A-Za-z0-9._-]+", "_", value).strip("._")
    return value[:100] or "tender"


def _clean_pdf_text(text: str) -> str:
    text = text.replace("\u00ad", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_requirement_section(text: str) -> str:
    """Return the most useful bid-document section, or a clear fallback."""
    match = _REQUIREMENT_HEADINGS.search(text)
    if not match:
        return "No clearly labelled bid-document requirement section was detected. Review the downloaded PDF."

    section = text[match.start():]
    end_match = _SECTION_END.search(section, match.end() - match.start())
    if end_match:
        section = section[:end_match.start()]
    section = _clean_pdf_text(section)
    return section[:12000]


def _download_pdf(url: str, destination: Path) -> bool:
    try:
        response = requests.get(
            url,
            timeout=45,
            headers={"User-Agent": "Amidel-eTender-Scraper/2.0"},
        )
        response.raise_for_status()
        if not response.content.startswith(b"%PDF"):
            logging.warning("Document link did not return a PDF: %s", url)
            return False
        destination.write_bytes(response.content)
        return True
    except Exception as exc:
        logging.warning("Could not download tender document %s: %s", url, exc)
        return False


def enrich_tender_documents(tender: dict, documents_root: str) -> dict:
    """Download all PDF links for one tender and add its requirements text."""
    links = tender.get("_DOCUMENT_LINKS") or []
    if not links and tender.get("LINK"):
        links = [tender["LINK"]]

    tender["REQUIREMENTS"] = ""
    tender["_DOCUMENT_LINKS"] = list(dict.fromkeys(links))
    if not links:
        tender["REQUIREMENTS"] = "No accompanying PDF links were found."
        return tender

    tender_dir = Path(documents_root) / _safe_filename(
        tender.get("TENDER_ID") or tender.get("TENDER_DESCRIPTION") or "tender"
    )
    tender_dir.mkdir(parents=True, exist_ok=True)
    sections = []

    for index, url in enumerate(links, 1):
        filename = _safe_filename(Path(urlparse(url).path).name) or f"document_{index}.pdf"
        if not filename.lower().endswith(".pdf"):
            filename += ".pdf"
        pdf_path = tender_dir / f"{index:02d}_{filename}"
        if not pdf_path.exists() and not _download_pdf(url, pdf_path):
            continue
        try:
            with fitz.open(pdf_path) as pdf:
                text = "\n".join(page.get_text() for page in pdf)
            sections.append(f"{filename}\n{_extract_requirement_section(text)}")
        except Exception as exc:
            logging.warning("Could not parse tender PDF %s: %s", pdf_path, exc)

    tender["REQUIREMENTS"] = "\n\n".join(sections) or "PDF links found, but no PDF could be downloaded or parsed."
    return tender


def create_tender_requirements_docx(tenders: Iterable[dict], batch_folder: str) -> str | None:
    """Create one Word report containing requirements for all eTender tenders."""
    tenders = list(tenders)
    if not tenders:
        return None

    output_path = Path(batch_folder) / "end product" / "Tender_Bid_Document_Requirements.docx"
    document = Document()
    document.add_heading("Tender Bid-Document Requirements", 0)
    document.add_paragraph(
        "Requirements extracted from accompanying eTender PDF documents. "
        "The original PDFs are retained in the batch documents folder for verification."
    )

    for index, tender in enumerate(tenders, 1):
        document.add_heading(
            f"{index}. {tender.get('TENDER_ID') or 'Tender'}", level=1
        )
        document.add_paragraph(tender.get("TENDER_DESCRIPTION") or "")
        document.add_paragraph(
            f"Department: {tender.get('DEPARTMENT') or 'Not stated'}\n"
            f"Closing date: {tender.get('CLOSING_DATE') or 'Not stated'}"
        )
        document.add_heading("Required bid documents", level=2)
        for paragraph in (tender.get("REQUIREMENTS") or "Not extracted.").split("\n"):
            paragraph = paragraph.strip()
            if paragraph:
                document.add_paragraph(paragraph, style="List Bullet")
        document.add_heading("Source PDFs", level=2)
        for url in tender.get("_DOCUMENT_LINKS") or []:
            document.add_paragraph(url, style="List Bullet")

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Aptos"
            run.font.size = Pt(10)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    logging.info("Tender requirements Word report created: %s", output_path)
    return str(output_path)